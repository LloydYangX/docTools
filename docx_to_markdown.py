#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
脚本功能: 将docx文档转换为markdown格式
依赖库: python-docx, markdown
使用方法: python docx_to_markdown.py 输入文件.docx [输出文件.md]
"""

import os
import sys
import re
import tempfile
import zipfile
import shutil
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def extract_images(docx_path, output_dir):
    """
    从docx文件中提取图片并保存到指定目录
    
    Args:
        docx_path (str): docx文件路径
        output_dir (str): 图片输出目录
    
    Returns:
        dict: 图片ID与输出路径的映射
    """
    # 确保输出目录存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 创建一个临时目录来解压docx文件
    temp_dir = tempfile.mkdtemp()
    
    try:
        # 将docx解压到临时目录
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # 检查media目录
        media_dir = os.path.join(temp_dir, 'word', 'media')
        if not os.path.exists(media_dir):
            return {}
        
        # 读取关系文件，映射图片ID
        rel_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        image_rels = {}
        
        if os.path.exists(rel_path):
            try:
                from lxml import etree
                tree = etree.parse(rel_path)
                root = tree.getroot()
                ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                
                for rel in root.xpath('//r:Relationship', namespaces=ns):
                    rid = rel.get('Id')
                    target = rel.get('Target')
                    if 'image' in rel.get('Type', '').lower() and target:
                        image_rels[rid] = target.split('/')[-1]
                        
            except ImportError:
                # 如果没有lxml库，就使用简单的正则表达式
                with open(rel_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    pattern = r'<Relationship Id="(rId\d+)" .*?Type=".*?image.*?" .*?Target=".*?/([^/]+)"'
                    for match in re.finditer(pattern, content):
                        rid, filename = match.groups()
                        image_rels[rid] = filename
        
        # 复制图片到输出目录并构建映射
        image_map = {}
        for filename in os.listdir(media_dir):
            src_path = os.path.join(media_dir, filename)
            if os.path.isfile(src_path):
                dst_path = os.path.join(output_dir, filename)
                shutil.copy2(src_path, dst_path)
                
                # 找到对应的rId
                for rid, rel_filename in image_rels.items():
                    if rel_filename == filename:
                        image_map[rid] = os.path.join('images', filename)
                        break
                else:
                    # 如果在关系文件中找不到，就使用文件名作为键
                    image_map[filename] = os.path.join('images', filename)
        
        return image_map
        
    finally:
        # 清理临时目录
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

def docx_to_markdown(input_path, output_path=None):
    """
    将docx文档转换为markdown格式
    
    Args:
        input_path (str): 输入docx文件路径
        output_path (str, optional): 输出markdown文件路径。如果未提供，将生成同名.md文件
    
    Returns:
        str: 输出文件的路径
    """
    if not os.path.exists(input_path):
        print(f"错误: 文件 '{input_path}' 不存在!")
        return None
        
    # 如果没有指定输出路径，则生成默认输出路径
    if output_path is None:
        file_dir = os.path.dirname(input_path)
        file_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(file_dir, file_name + ".md")
    
    # 创建图片目录
    img_dir = os.path.join(os.path.dirname(output_path), 'images')
    if not os.path.exists(img_dir):
        os.makedirs(img_dir)
    
    try:
        # 提取图片
        image_map = extract_images(input_path, img_dir)
        
        # 加载文档
        doc = Document(input_path)
        
        # 准备markdown内容
        markdown_content = []
        
        # 处理文档标题（如果有）
        if len(doc.paragraphs) > 0 and doc.paragraphs[0].style.name.startswith('Title'):
            title_text = doc.paragraphs[0].text.strip()
            if title_text:
                markdown_content.append(f"# {title_text}\n")
        
        # 处理段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                markdown_content.append("")
                continue
            
            # 处理标题
            style_name = para.style.name.lower()
            if 'heading' in style_name:
                heading_level = 0
                if '1' in style_name:
                    heading_level = 1
                elif '2' in style_name:
                    heading_level = 2
                elif '3' in style_name:
                    heading_level = 3
                elif '4' in style_name:
                    heading_level = 4
                else:
                    heading_level = 5  # 默认最多5级标题
                
                if heading_level > 0:
                    markdown_content.append(f"{'#' * heading_level} {text}\n")
                    continue
            
            # 检查段落格式
            is_bullet = False
            is_numbered = False
            
            # 检查段落属性或样式来确定是否为列表
            if hasattr(para, '_p') and para._p is not None:
                if para._p.pPr is not None and para._p.pPr.numPr is not None:
                    if para._p.pPr.numPr.numId is not None:
                        num_id = para._p.pPr.numPr.numId.val
                        if num_id != 0:
                            # 检查是否为有序列表或无序列表
                            is_numbered = True  # 默认认为是有序列表
                            is_bullet = 'bullet' in style_name or 'list' in style_name
            
            # 处理列表
            if is_bullet:
                markdown_content.append(f"* {text}")
            elif is_numbered:
                markdown_content.append(f"1. {text}")  # Markdown会自动处理编号
            else:
                # 处理段落中的图片
                img_refs = []
                for run in para.runs:
                    for element in run._element.findall('.//a:graphic', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                        # 尝试找到图片ID
                        for rid_elem in element.findall('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                            rid = rid_elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if rid and rid in image_map:
                                img_refs.append(f"![图片]({image_map[rid]})")
                
                # 如果段落中有图片，添加图片引用
                if img_refs:
                    markdown_content.append(text)
                    markdown_content.extend(img_refs)
                else:
                    markdown_content.append(text)
        
        # 处理表格
        for table in doc.tables:
            # 添加表格的markdown格式
            table_markdown = []
            
            # 处理表头
            header_row = []
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    header_row.append(cell.text.strip() or " ")
                table_markdown.append("| " + " | ".join(header_row) + " |")
                
                # 添加分隔行
                separator = ["---" for _ in range(len(header_row))]
                table_markdown.append("| " + " | ".join(separator) + " |")
                
                # 处理数据行
                for i, row in enumerate(table.rows):
                    if i == 0:  # 跳过标题行
                        continue
                    
                    data_row = []
                    for cell in row.cells:
                        data_row.append(cell.text.strip() or " ")
                    
                    table_markdown.append("| " + " | ".join(data_row) + " |")
                
                # 添加空行
                table_markdown.append("")
            
            markdown_content.append("\n".join(table_markdown))
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(markdown_content))
        
        print(f"已成功将docx文档转换为markdown格式")
        print(f"输出文件: {output_path}")
        
        if image_map:
            print(f"已提取 {len(image_map)} 张图片到 {img_dir} 目录")
        
        return output_path
        
    except Exception as e:
        print(f"转换过程中出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数，处理命令行参数并执行转换操作"""
    if len(sys.argv) < 2:
        print("用法: python docx_to_markdown.py 输入文件.docx [输出文件.md]")
        return
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    docx_to_markdown(input_path, output_path)

if __name__ == "__main__":
    main() 