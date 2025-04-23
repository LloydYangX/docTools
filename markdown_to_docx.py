#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
脚本功能: 将Markdown文档转换为Word文档(.docx)格式
依赖库: markdown, python-docx, beautifulsoup4, requests
使用方法: python markdown_to_docx.py 输入文件.md [输出文件.docx]
"""

import os
import sys
import re
import shutil
import tempfile
from pathlib import Path
import base64
from urllib.parse import urlparse
import requests
from bs4 import BeautifulSoup

try:
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("请安装所需依赖：pip install Markdown python-docx beautifulsoup4 requests")
    sys.exit(1)

def extract_images_from_markdown(md_text, base_dir):
    """
    从Markdown文本中提取图片链接，下载图片并返回替换后的文本和图片映射
    
    Args:
        md_text (str): Markdown文本内容
        base_dir (str): Markdown文件所在目录
    
    Returns:
        tuple: (替换后的Markdown文本, 图片映射字典)
    """
    image_pattern = r'!\[(.*?)\]\((.*?)\)'
    images = {}
    img_counter = 0
    
    def replace_image(match):
        nonlocal img_counter
        img_counter += 1
        alt_text = match.group(1)
        img_url = match.group(2)
        
        # 处理本地图片
        if not img_url.startswith(('http://', 'https://', 'data:')):
            if os.path.isabs(img_url):
                img_path = img_url
            else:
                # 相对路径
                img_path = os.path.join(base_dir, img_url)
                
            if os.path.exists(img_path):
                img_filename = f"image_{img_counter}{os.path.splitext(img_path)[1]}"
                images[img_filename] = (img_path, alt_text)
                return f"![{alt_text}]({img_filename})"
        
        # 处理网络图片或已经替换的图片
        elif img_url.startswith(('http://', 'https://')):
            try:
                img_filename = f"image_{img_counter}.jpg"  # 默认扩展名
                parsed_url = urlparse(img_url)
                if parsed_url.path:
                    ext = os.path.splitext(parsed_url.path)[1]
                    if ext:
                        img_filename = f"image_{img_counter}{ext}"
                
                images[img_filename] = (img_url, alt_text)
                return f"![{alt_text}]({img_filename})"
            except Exception as e:
                print(f"警告：无法处理图片URL：{img_url}, 错误：{str(e)}")
        
        # 返回原始匹配，不做修改
        return match.group(0)
    
    # 替换并提取图片
    processed_text = re.sub(image_pattern, replace_image, md_text)
    return processed_text, images

def download_images(images, output_dir):
    """
    下载或复制图片到临时目录
    
    Args:
        images (dict): 图片映射字典
        output_dir (str): 图片保存目录
    
    Returns:
        dict: 更新后的图片映射，包含本地路径
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    image_paths = {}
    for filename, (url_or_path, alt) in images.items():
        dest_path = os.path.join(output_dir, filename)
        
        try:
            if url_or_path.startswith(('http://', 'https://')):
                # 下载网络图片
                response = requests.get(url_or_path, stream=True)
                if response.status_code == 200:
                    with open(dest_path, 'wb') as f:
                        response.raw.decode_content = True
                        shutil.copyfileobj(response.raw, f)
                    image_paths[filename] = (dest_path, alt)
                else:
                    print(f"警告：无法下载图片 {url_or_path}, 状态码：{response.status_code}")
                    continue
            else:
                # 复制本地图片
                shutil.copy2(url_or_path, dest_path)
                image_paths[filename] = (dest_path, alt)
        except Exception as e:
            print(f"警告：处理图片失败 {url_or_path}, 错误：{str(e)}")
    
    return image_paths

def convert_html_to_docx(html, image_paths, output_path):
    """
    将HTML转换为DOCX格式
    
    Args:
        html (str): HTML内容
        image_paths (dict): 图片路径字典
        output_path (str): 输出DOCX文件路径
    """
    # 创建新文档
    doc = Document()
    
    # 配置文档基础样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 解析HTML
    soup = BeautifulSoup(html, 'html.parser')
    
    # 处理元素
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'pre', 'blockquote', 'img', 'table']):
        if element.name.startswith('h') and len(element.name) == 2:
            # 处理标题
            level = int(element.name[1])
            heading = doc.add_heading(level=level)
            heading.add_run(element.get_text())
        
        elif element.name == 'p':
            # 处理段落
            p = doc.add_paragraph()
            
            # 遍历段落内容
            for content in element.contents:
                if content.name == 'strong' or content.name == 'b':
                    p.add_run(content.get_text()).bold = True
                elif content.name == 'em' or content.name == 'i':
                    p.add_run(content.get_text()).italic = True
                elif content.name == 'a':
                    run = p.add_run(content.get_text())
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
                elif content.name == 'code':
                    run = p.add_run(content.get_text())
                    run.font.name = 'Courier New'
                elif content.name == 'img':
                    # 图片会单独处理
                    pass
                else:
                    # 处理纯文本
                    if isinstance(content, str):
                        p.add_run(content)
                    else:
                        p.add_run(content.get_text())
        
        elif element.name == 'img':
            # 处理图片
            src = element.get('src')
            alt = element.get('alt', '')
            
            if src in image_paths:
                img_path, _ = image_paths[src]
                try:
                    doc.add_picture(img_path, width=Inches(6))  # 默认宽度
                    # 添加图片说明
                    if alt:
                        caption = doc.add_paragraph(alt)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.style = 'Caption'
                except Exception as e:
                    print(f"警告：无法插入图片 {src}, 错误：{str(e)}")
        
        elif element.name == 'ul' or element.name == 'ol':
            # 处理列表
            for li in element.find_all('li', recursive=False):
                level = 0
                parent = li.parent
                while parent and parent.name in ['ul', 'ol']:
                    level += 1
                    parent = parent.parent
                
                # 设置缩进级别
                p = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                p.paragraph_format.left_indent = Pt(18 * level)
                p.add_run(li.get_text())
        
        elif element.name == 'pre':
            # 处理代码块
            code = element.get_text()
            p = doc.add_paragraph(style='No Spacing')
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        elif element.name == 'blockquote':
            # 处理引用
            p = doc.add_paragraph(style='Intense Quote')
            p.add_run(element.get_text())
        
        elif element.name == 'table':
            # 处理表格
            rows = len(element.find_all('tr'))
            if rows > 0:
                first_row = element.find('tr')
                cols = len(first_row.find_all(['th', 'td']))
                
                if cols > 0 and rows > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(element.find_all('tr')):
                        cells = row.find_all(['th', 'td'])
                        for j, cell in enumerate(cells):
                            if j < cols:  # 防止表格不规则
                                table.cell(i, j).text = cell.get_text()
    
    # 保存文档
    doc.save(output_path)

def markdown_to_docx(input_path, output_path=None):
    """
    将Markdown文档转换为Word文档(.docx)格式
    
    Args:
        input_path (str): 输入Markdown文件路径
        output_path (str, optional): 输出Word文件路径。如果未提供，将生成同名.docx文件
    
    Returns:
        str: 输出文件的路径
    """
    if not os.path.exists(input_path):
        print(f"错误: 文件 '{input_path}' 不存在!")
        return None
    
    # 如果没有指定输出路径，则生成默认输出路径
    if output_path is None:
        file_dir = os.path.dirname(input_path)
        file_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(file_dir, file_name + ".docx")
    
    try:
        # 读取Markdown文件
        with open(input_path, 'r', encoding='utf-8') as f:
            md_text = f.read()
        
        # 提取并替换图片链接
        base_dir = os.path.dirname(os.path.abspath(input_path))
        processed_md, images = extract_images_from_markdown(md_text, base_dir)
        
        # 创建临时目录存放图片
        temp_dir = tempfile.mkdtemp()
        try:
            # 下载或复制图片
            img_dir = os.path.join(temp_dir, 'images')
            image_paths = download_images(images, img_dir)
            
            # 转换Markdown为HTML
            html = markdown.markdown(
                processed_md,
                extensions=[
                    'markdown.extensions.tables',
                    'markdown.extensions.fenced_code',
                    'markdown.extensions.nl2br',
                    'markdown.extensions.sane_lists'
                ]
            )
            
            # 将HTML转换为DOCX
            convert_html_to_docx(html, image_paths, output_path)
            
            print(f"已成功将Markdown文档转换为Word格式")
            print(f"输出文件: {output_path}")
            return output_path
            
        finally:
            # 清理临时目录
            shutil.rmtree(temp_dir)
            
    except Exception as e:
        print(f"转换过程中出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数，处理命令行参数并执行转换操作"""
    if len(sys.argv) < 2:
        print("用法: python markdown_to_docx.py 输入文件.md [输出文件.docx]")
        return
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    markdown_to_docx(input_path, output_path)

if __name__ == "__main__":
    main() 