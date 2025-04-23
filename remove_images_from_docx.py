#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
脚本功能: 解析docx文档，移除所有图片，然后生成一个新的不含图片的docx文档
依赖库: python-docx
使用方法: python remove_images_from_docx.py 输入文件.docx [输出文件.docx]
"""

import os
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

def remove_images_from_document(input_path, output_path=None):
    """
    从Word文档中移除所有图片，并保存为新文档
    
    Args:
        input_path (str): 输入docx文件路径
        output_path (str, optional): 输出docx文件路径。如果未提供，将生成'无图片_' + 原文件名
    
    Returns:
        str: 输出文件的路径
    """
    if not os.path.exists(input_path):
        print(f"错误: 文件 '{input_path}' 不存在!")
        return None
        
    # 如果没有指定输出路径，则生成默认输出路径
    if output_path is None:
        file_dir = os.path.dirname(input_path)
        file_name = os.path.basename(input_path)
        output_path = os.path.join(file_dir, "无图片_" + file_name)
    
    try:
        # 加载文档
        doc = Document(input_path)
        
        # 查找并移除所有图片
        removed_count = 0
        
        # 处理文档中的内联图片
        for para in doc.paragraphs:
            # 遍历段落中的所有运行(run)
            runs_to_remove = []
            for i, run in enumerate(para.runs):
                # 检查运行中是否有图片元素
                if run._element.findall('.//a:graphic', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                    runs_to_remove.append(i)
                    removed_count += 1
            
            # 从后向前删除，避免索引变化
            for idx in sorted(runs_to_remove, reverse=True):
                para._p.remove(para.runs[idx]._element)
        
        # 处理表格中的图片
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        runs_to_remove = []
                        for i, run in enumerate(para.runs):
                            if run._element.findall('.//a:graphic', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                                runs_to_remove.append(i)
                                removed_count += 1
                        
                        # 从后向前删除
                        for idx in sorted(runs_to_remove, reverse=True):
                            para._p.remove(para.runs[idx]._element)
        
        # 处理文档中的浮动图片和图形
        for element in doc._element.xpath('//w:drawing | //w:pict'):
            element.getparent().remove(element)
            removed_count += 1
        
        # 保存新文档
        doc.save(output_path)
        print(f"成功处理文档，已移除 {removed_count} 个图片/图形元素")
        print(f"新文档已保存至: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"处理文档时出错: {str(e)}")
        return None

def main():
    """主函数，处理命令行参数并执行图片移除操作"""
    if len(sys.argv) < 2:
        print("用法: python remove_images_from_docx.py 输入文件.docx [输出文件.docx]")
        return
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    remove_images_from_document(input_path, output_path)

if __name__ == "__main__":
    main() 